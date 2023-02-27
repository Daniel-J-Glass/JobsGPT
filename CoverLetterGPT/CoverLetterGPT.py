#!/usr/bin/env python3

import sys
import os
import json

import re
from unidecode import unidecode
import time

import logging

import docx

import ChatGPT

def replace_weird_unicode(text):
    return unidecode(text)

class CoverLetterGPT:
    #TODO: Store cover letters for similar jobs (replace job name) and reuse them to decrease API usage
    gpt = ChatGPT.ChatGPT("config.json")
    summarizer = ChatGPT.ChatGPT("summarizer_config.json")

    preprompt = "The following cover letter has been tailored specifically for the job role and description, using the most impressive details from your resume. It is a highly engaging, accurately written, and impressive cover letter that has been designed to showcase your unique strengths and qualifications, and to make you stand out as a top candidate. The letter incorporates a mix of sentence lengths, a natural rhythm, and a good amount of compound sentences to create a smooth and professional tone. The format and spacing have been optimized for readability, with concise and compelling content that emphasizes your skills and qualifications in accordance with the job description. The cover letter does not include years of experience or skills that are not represented on the resume, but rather focuses on the key strengths that make you the ideal candidate for the job. It is guaranteed to meet your expectations and create an outstanding first impression with the hiring manager. The following is a cover letter that meets these criteria and begins with \"Dear Hiring Manager\":"
    prompt_layout = "Resume Details: {resume}\n\nJob Summary: {job_summary}\n\n{preprompt}"
    chat_layout = "I am a highly intelligent question answering bot. If you ask me a question that is rooted in truth, I will give you the answer.\n\
        Resume Details: {resume}\n\n{preprompt}"


    resume_summ_prompt = "Resume: {resume}\n\nHere is a complete breakdown of the individual's job experience, including job details and dates, as well as the impacts they've had on their respective companies. Their technical experience and skills are also included, with no details left out. Personal details, such as their first and last name, are also provided. The information is organized by employment dates and presented in an easy-to-read bullet point format."
    job_summ_prompt = "Job Description: {document}\n\nHere are full notes of the job description; position name, job role, the company, a comprehensive list of technical skills and experience are included. Using bullet points:"

    cover_letter_check = "Evaluate the following cover letter to determine if it is suitable to send to a hiring manager. Responses should be 'YES' if the cover letter meets the criteria and 'NO' followed by the reason if it does not. The cover letter must be natural-sounding and avoid generic language or filler values like 'Company ABC'.\n\n{cover_letter}"
    def __init__(self,job_description, resume_path, extra_details = None, gpt=gpt, preprompt = preprompt, prompt_layout = prompt_layout, cover_letter_check = cover_letter_check, checker_retries = 2):
        self.job_description = job_description
        self.resume = self._docx_to_summary(resume_path) # TODO: Parse resume into plain text
        self.extra_details = extra_details

        self.job_summary = self._summarize(self.job_summ_prompt, self.job_description)
        
        self.preprompt=preprompt
        self.prompt_layout = prompt_layout

        self.cover_letter_check = cover_letter_check
        self.checker_retries = checker_retries

        self.gpt = gpt

        # Add a file handler to the logger
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.INFO)
        log_file_path = "cover_letter_gpt.log"
        handler = logging.FileHandler(log_file_path)
        formatter = logging.Formatter('%(asctime)s - %(message)s')
        handler.setFormatter(formatter)
        self.logger.addHandler(handler)

    def _summarize(self,prompt,doc):
        """
        Generates summary from document using AI

        Args:
            prompt (str): The prompt for summarizing the specific type of document
            doc (str): The input document to be summarized
        """
        summary = self.summarizer.chat(prompt.format(document=doc))
        summary = replace_weird_unicode(summary)
        return summary

    def _docx_to_summary(self, docx_file_path):
        """
        Converts a .docx file to a summary.

        Args:
            docx_file_path (str): The path to the .docx file to be summarized.

        Returns:
            str: The summary of the .docx file as a string.
        """
        doc_name = os.path.splitext(os.path.basename(docx_file_path))[0]
        summary_file_path = "summaries.json"
        summaries = {}

        # Check if the summary file already exists
        if os.path.exists(summary_file_path):
            with open(summary_file_path) as f:
                summaries = json.load(f)

        # Check if a summary already exists for this document
        if doc_name in summaries:
            summary = summaries[doc_name]
        else:
            # Open the .docx file
            docx_file = docx.Document(docx_file_path)

            # Read the text from the .docx file
            text = "\n".join([paragraph.text for paragraph in docx_file.paragraphs])

            # Generate a summary of the text with larger model
            response = self.gpt.chat(self.resume_summ_prompt.format(resume=text))
            summary = response.strip()

            # Update the summaries dictionary with the new summary
            summaries[doc_name] = summary

            # Save the updated summaries to the summary file
            with open(summary_file_path, "w") as f:
                json.dump(summaries, f)

        return summary

    def _str_to_docx(self, text, docx_file_path):
        """
        Converts a Python string to a .docx file.

        Args:
        text (str): the string to be converted to a .docx file.
        docx_file_path (str): the desired path for the output .docx file.

        Returns:
        None
        """
        # Create a new .docx file
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'

        # Add the text to a new paragraph in the document
        p = doc.add_paragraph(text)

        # Save the .docx file
        doc.save(docx_file_path)

    def _cover_letter_is_valid(self, cover_letter):
        prompt = self.cover_letter_check.format(cover_letter=cover_letter)

        # Recieve quality check
        response = self.gpt.chat(prompt)

        self.logger.info(response)

        return "YES" in response

    def _generate_cover_letter_text(self, prompt):
        # Recieve full prompt
        cover_letter = self.gpt.chat(prompt)

        cover_letter = replace_weird_unicode(cover_letter)

        self.logger.info(f"Generated cover letter: {cover_letter}")

        return cover_letter
    
    def ask(self,prompt):
        prompt = self.chat_layout.format(preprompt=prompt, resume = self.resume)
        return self.summarizer.chat(prompt)

    def generate_cover_letter(self, output_path = "CoverLetter.docx",extra_prompting = None):
        # Adding extra info
        preprompt = self.preprompt
        if extra_prompting:
            preprompt+=f"\n{extra_prompting}"

        # Creating full prompt
        prompt = self.prompt_layout.format(preprompt=self.preprompt, resume = self.resume, job_summary = self.job_summary)

        if self.extra_details:
            prompt+= f"\n{self.extra_details}"

        # measuring response time
        start_time = time.time()
    
        # Ensuring cover letter validity
        is_valid = False
        for i in range(self.checker_retries):
            if not is_valid:
                cover_letter = self._generate_cover_letter_text(prompt)
                is_valid = self._cover_letter_is_valid(cover_letter)
            
            if not is_valid:
                continue

            self._str_to_docx(cover_letter,output_path)
            elapsed_time = time.time() - start_time  # Calculate the elapsed time
            self.logger.info(f"Cover letter generation took {elapsed_time} seconds and {i} retries")
            return
        
        raise ValueError("ChatGPT was unable to adequately generate a cover letter.")

def main():
    job_description = \
    r'''
    Ethisphere is seeking an intellectually curious Data Scientist to gather, analyze and compile data needed to identify trends and patterns across data sets. The Data Scientist will partner with technology teams to inform and build long-term data strategy. 

About Us 

At Ethisphere, we believe that strong ethics is good business. We help companies advance business integrity through data-driven assessments, benchmarking, and guidance. We honor superior achievement through the World’s Most Ethical Companies® recognition, we promote professional collaboration through the Business Ethics Leadership Alliance (BELA), and we showcase best practices through events and media such as the Global Ethics Summit and Ethisphere Magazine. 

Our employees are our greatest strength. We seek the brightest talent so we may invest in their professional development and position them for success within the expanding ethics, compliance, and corporate culture space. Our fully remote team consists of more than 55 employees living in 20 states across the U.S. and four time zones. We have a welcoming, supportive, and engaged culture that values flexibility, collaboration, and strong work-life balance. With a history of strong, sustained growth, we offer professional opportunities at junior, middle, and senior levels.

Ethisphere has a strategic partnership with predictis, a business platform that partners with world-class data software businesses. The platform is part of Alpine Investors, a people-driven private equity firm.

Responsibilities 

Build out long-term work plans while also managing day-today priorities and initiatives
Develop a detailed understanding of Ethisphere’s products and data
Collaborate with SME to define questions to evaluate, analyze data and present findings
Evaluate data inputs and provide recommendations for data transformations and storage
Organize and manipulate data so that it can be analyzed
Develop algorithms which process data to generate desired insight 
Write complex queries to extract and analyze data
Develop an approach for rapid prototype reports
Develop a maintainable reporting framework 
Work collaboratively with leadership and SMEs to provide high-impact insights 
Build reports and dashboards in data analytics tools
Partner with software engineering to prototype new reporting concepts
Qualifications 

3+ years in a data science, data analytics role
3+ using data analytics tools such as Power BI or Tableau
Bachelor’s Degree or higher in a quantitative discipline
Excellent analytical skills - the ability to identify trends, patterns and insights from data
Experience working with relational databases (MySQL, SQL) and data warehousing
Experience with data querying languages, statistical or mathematical software.
Ability to work at a conceptual and detail level
Experience working with BI, data visualization, or data statistical platforms such as Tableau, Power BI, SAS, SPSS, Excel, or similar
Experience working with SMEs to get tangible requirements and domain understanding
Demonstrated ability to analyze and extract insights from large sets of data
Experience utilizing and connecting information across multiple sources to produce reports and insights
    '''
    resume_path = "C:/Users/drago/Documents/Programming/DOWORK/AutoApply/Daniel Glass Resume - Data Science.docx"
    cover_gen = CoverLetterGPT(job_description=job_description,resume_path=resume_path)
    cover_gen.generate_cover_letter("Daniel Glass - CoverLetter.docx")

def test():
    gpt = CoverLetterGPT("","C:/Users/drago/Documents/Programming/DOWORK/AutoApply/Daniel Glass Resume - Data Science.docx")
    print(gpt._generate_cover_letter_text(prompt="output 5 apostrophes"))

if __name__ =="__main__":
    main()